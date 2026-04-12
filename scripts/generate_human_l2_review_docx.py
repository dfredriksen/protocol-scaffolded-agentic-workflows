from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "hitl" / "reviewer-packet-docx"

DOCS = [
    ("docs/hitl/H005-l2-reviewer-assignment-packet.md", "H005-l2-reviewer-assignment-packet.docx"),
    ("docs/hitl/H007-l2-review-kickoff-note.md", "H007-l2-review-kickoff-note.docx"),
    ("docs/hitl/H008-l2-reviewer-checklist-v1.md", "H008-l2-reviewer-checklist-v1.docx"),
    ("docs/hitl/H009-drive-facing-l2-review-handoff-v1.md", "H009-human-l2-review-handoff-v1.docx"),
    ("docs/hitl/H011-reviewer-readme-v1.md", "H011-reviewer-readme-v1.docx"),
    ("docs/hitl/H012-human-l2-review-attachment-index-v1.md", "H012-human-l2-review-attachment-index-v1.docx"),
    ("docs/hitl/H014-human-l2-review-packet-manifest-v1.md", "H014-human-l2-review-packet-manifest-v1.docx"),
    ("docs/benchmarks/l2-hosted-execution-bundle-v1.md", "l2-hosted-execution-bundle-v1.docx"),
    ("docs/process/l2-review-handoff-protocol.md", "l2-review-handoff-protocol.docx"),
    ("docs/benchmarks/scoring-rubric-v1.md", "scoring-rubric-v1.docx"),
    ("docs/benchmarks/review-pass-worksheet-template.md", "review-pass-worksheet-template.docx"),
    ("docs/benchmarks/blinded-review-worksheet-template.md", "blinded-review-worksheet-template.docx"),
    ("docs/benchmarks/l2-hosted-review-summary-template.md", "l2-hosted-review-summary-template.docx"),
]


def ensure_code_style(document: Document) -> str:
    styles = document.styles
    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(9)
    return "CodeBlock"


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | tuple[None, int]:
    if start + 1 >= len(lines):
        return None, start
    first = lines[start].strip()
    second = lines[start + 1].strip()
    if not (first.startswith("|") and first.endswith("|")):
        return None, start
    if not re.fullmatch(r"\|(?:\s*:?-+:?\s*\|)+", second):
        return None, start
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines):
        line = lines[idx].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        row = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(row)
        idx += 1
    return rows, idx


def render_markdown(source: Path, destination: Path) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    code_style = ensure_code_style(doc)

    idx = 0
    in_code = False
    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            idx += 1
            continue

        if in_code:
            p = doc.add_paragraph(style=code_style)
            p.add_run(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        table_rows, next_idx = parse_table(lines, idx)
        if table_rows:
            table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(table_rows):
                for c, cell in enumerate(row):
                    table.cell(r, c).text = cell
            idx = next_idx
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            heading_text = stripped[level:].strip()
            doc.add_heading(heading_text, level=level - 1)
            idx += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, numbered.group(2))
            idx += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            idx += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        idx += 1

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> None:
    generated: list[str] = []
    for rel_src, out_name in DOCS:
        src = ROOT / rel_src
        dst = OUTPUT_DIR / out_name
        render_markdown(src, dst)
        generated.append(out_name)

    manifest = OUTPUT_DIR / "README.txt"
    manifest.write_text(
        "Human L2 reviewer packet .docx bundle\n\n"
        + "\n".join(generated)
        + "\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
