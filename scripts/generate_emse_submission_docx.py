from __future__ import annotations

import re
import shutil
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table as PdfTable, TableStyle


ROOT = Path(__file__).resolve().parents[1]
BUNDLE_DIR = ROOT / "docs" / "evidence" / "review-packets" / "emse-submission-v1"
UPLOAD_DIR = BUNDLE_DIR / "upload-set"
BRAND_LOGO = Path("C:/Users/Dan/Desktop/Projects/Quantyra/docs/assets/branding/logo_text.png")
BRAND_COLOR = RGBColor(0x56, 0x3C, 0x99)

REVIEW_DOCS = [
    ("docs/evidence/manuscript-draft-v1.md", "manuscript-draft-v1.docx"),
    ("docs/evidence/emse-cover-letter-draft-v1.md", "emse-cover-letter-draft-v1.docx"),
    ("docs/evidence/manuscript-supplement-v1.md", "manuscript-supplement-v1.docx"),
    ("docs/evidence/emse-title-page-draft-v1.md", "emse-title-page-draft-v1.docx"),
    ("docs/evidence/post-l2-journal-readiness-status-v1.md", "post-l2-journal-readiness-status-v1.docx"),
]

UPLOAD_DOCS = [
    ("docs/evidence/manuscript-draft-v1.md", "EMSE_main_manuscript.docx"),
    ("docs/evidence/emse-title-page-draft-v1.md", "EMSE_title_page_and_declarations.docx"),
    ("docs/evidence/emse-cover-letter-draft-v1.md", "EMSE_cover_letter.docx"),
    ("docs/evidence/manuscript-supplement-v1.md", "EMSE_supplement_source.docx"),
]

SYNC_FILES = [
    "docs/evidence/manuscript-draft-v1.md",
    "docs/evidence/manuscript-claim-matrix-v1.md",
    "docs/evidence/manuscript-tables-v1.md",
    "docs/evidence/figure-captions-draft-v1.md",
    "docs/evidence/emse-cover-letter-draft-v1.md",
    "docs/evidence/emse-title-page-draft-v1.md",
    "docs/evidence/emse-title-page-finalization-checklist-v1.md",
    "docs/evidence/manuscript-supplement-v1.md",
    "docs/evidence/manuscript-packet-v1.md",
    "docs/evidence/post-l2-journal-readiness-status-v1.md",
    "docs/evidence/journal-submission-packet-checklist-v1.md",
    "docs/evidence/emse-submission-preparation-note-v1.md",
    "docs/evidence/emse-production-readiness-rubric-v1.md",
    "docs/evidence/disclosure-clearance-note-v1.md",
    "docs/evidence/submission-signoff-record-v1.md",
]


def ensure_code_style(document: Document) -> str:
    styles = document.styles
    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(9)
    return "CodeBlock"


def ensure_caption_style(document: Document) -> str:
    styles = document.styles
    if "TableCaption" not in styles:
        style = styles.add_style("TableCaption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True
    return "TableCaption"


def ensure_reference_style(document: Document) -> str:
    styles = document.styles
    if "ReferenceEntry" not in styles:
        style = styles.add_style("ReferenceEntry", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.3)
        style.paragraph_format.first_line_indent = Inches(-0.3)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
    return "ReferenceEntry"


def apply_base_styles(document: Document, branded: bool) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Arial" if branded else "Times New Roman"
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        if branded:
            style.font.color.rgb = BRAND_COLOR


def add_branding(document: Document, title: str) -> None:
    if BRAND_LOGO.exists():
        document.add_picture(str(BRAND_LOGO), width=Inches(3.2))
    title_para = document.add_paragraph()
    run = title_para.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_COLOR
    document.add_paragraph()


def add_plain_title(document: Document, title: str) -> None:
    para = document.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    document.add_paragraph()


def add_footer(section, footer_text: str | None) -> None:
    if not footer_text:
        return
    footer = section.footer.paragraphs[0]
    footer.text = footer_text
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    if footer_text == "Quantyra submission draft":
        footer.runs[0].font.color.rgb = BRAND_COLOR


def add_page_number_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = 1
    run = footer.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = " PAGE "

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_separate.append(text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*].+?\*\*|\*[^*].+?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue
        if part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            continue
        paragraph.add_run(part)


def set_cell_width(cell, width_inches: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))


def set_fixed_layout(table) -> None:
    tblPr = table._tbl.tblPr
    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def estimate_col_widths(col_count: int) -> list[float]:
    presets = {
        9: [0.75, 1.05, 0.52, 0.52, 0.52, 0.52, 0.52, 0.52, 1.48],
        8: [0.48, 0.95, 1.02, 0.52, 0.68, 0.68, 0.68, 1.39],
        4: [1.15, 1.65, 2.05, 2.05],
        3: [1.45, 2.25, 3.20],
    }
    if col_count in presets:
        return presets[col_count]
    width = 6.9 / col_count
    return [width] * col_count


def format_table(table, rows: list[list[str]], *, upload_manuscript: bool) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_fixed_layout(table)

    widths = estimate_col_widths(len(rows[0]))
    for col_idx, width in enumerate(widths):
        for row in table.rows:
            set_cell_width(row.cells[col_idx], width)

    set_repeat_table_header(table.rows[0])

    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                text = paragraph.text.strip()
                if upload_manuscript and (re.fullmatch(r"[0-9]+/[0-9]+", text) or re.fullmatch(r"[PWT][0-9]+", text)):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif upload_manuscript and len(text) <= 12 and " " not in text and c > 1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman" if run.font.name != "Consolas" else "Consolas"
                    run.font.size = Pt(8.5 if upload_manuscript else 9)
                    if r == 0:
                        run.bold = True
                        run.font.size = Pt(8.5 if upload_manuscript else 9)

    if upload_manuscript:
        for cell in table.rows[0].cells:
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "D9D9D9")
            cell._tc.get_or_add_tcPr().append(shading_elm)


def configure_sections(document: Document, *, upload_manuscript: bool) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        if upload_manuscript:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        else:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | tuple[None, int]:
    if start + 1 >= len(lines):
        return None, start
    first = lines[start].strip()
    second = lines[start + 1].strip()
    if not (first.startswith("|") and first.endswith("|")):
        return None, start
    if not re.fullmatch(r"\|(?:\s*:?-+:?\s*\|)+", second):
        return None, start
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines):
        line = lines[idx].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        rows.append([cell.strip() for cell in line.strip("|").split("|")])
        idx += 1
    return rows, idx


def render_markdown(
    source: Path,
    destination: Path,
    *,
    branded: bool,
    footer_text: str | None,
    include_page_numbers: bool,
) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    apply_base_styles(doc, branded=branded)
    code_style = ensure_code_style(doc)
    caption_style = ensure_caption_style(doc)
    reference_style = ensure_reference_style(doc)
    upload_manuscript = destination.name == "EMSE_main_manuscript.docx"
    title = lines[0].lstrip("#").strip() if lines and lines[0].startswith("#") else source.stem
    if branded:
        add_branding(doc, title)
    else:
        add_plain_title(doc, title)

    idx = 0
    in_code = False
    in_references = False
    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        stripped = line.strip()

        if idx == 0 and stripped.startswith("#"):
            idx += 1
            continue

        if stripped.startswith("```"):
            in_code = not in_code
            idx += 1
            continue

        if in_code:
            p = doc.add_paragraph(style=code_style)
            p.add_run(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        table_rows, next_idx = parse_table(lines, idx)
        if table_rows:
            table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            for r, row in enumerate(table_rows):
                for c, cell in enumerate(row):
                    paragraph = table.cell(r, c).paragraphs[0]
                    add_inline_runs(paragraph, cell)
            format_table(table, table_rows, upload_manuscript=upload_manuscript)
            idx = next_idx
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            heading_text = stripped[level:].strip()
            doc.add_heading(heading_text, level=level - 1)
            in_references = heading_text == "References"
            idx += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, numbered.group(2))
            idx += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            idx += 1
            continue

        paragraph_style = None
        if stripped.startswith("Table "):
            paragraph_style = caption_style
        elif in_references:
            paragraph_style = reference_style
        p = doc.add_paragraph(style=paragraph_style)
        add_inline_runs(p, stripped)
        idx += 1

    configure_sections(doc, upload_manuscript=upload_manuscript)

    for section in doc.sections:
        if include_page_numbers:
            add_page_number_footer(section)
        else:
            add_footer(section, footer_text)
        section.start_type = WD_SECTION.CONTINUOUS

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def strip_inline_code(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def rich_text(text: str) -> str:
    parts = re.split(r"(`[^`]+`|\*\*[^*].+?\*\*|\*[^*].+?\*)", text)
    rendered: list[str] = []
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            rendered.append(f"<font name='Courier'>{escape(part[1:-1])}</font>")
        elif part.startswith("**") and part.endswith("**"):
            rendered.append(f"<b>{escape(part[2:-2])}</b>")
        elif part.startswith("*") and part.endswith("*"):
            rendered.append(f"<i>{escape(part[1:-1])}</i>")
        else:
            rendered.append(escape(part))
    return "".join(rendered)


def render_markdown_pdf(source: Path, destination: Path) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].lstrip("#").strip() if lines and lines[0].startswith("#") else source.stem

    stylesheet = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=stylesheet["BodyText"], fontName="Times-Roman", fontSize=11, leading=14, spaceAfter=8)
    code = ParagraphStyle("Code", parent=stylesheet["BodyText"], fontName="Courier", fontSize=8.5, leading=10, spaceAfter=6)
    h1 = ParagraphStyle("H1", parent=stylesheet["Heading1"], fontName="Times-Bold", fontSize=16, leading=20, spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=stylesheet["Heading2"], fontName="Times-Bold", fontSize=13, leading=16, spaceAfter=8)
    h3 = ParagraphStyle("H3", parent=stylesheet["Heading3"], fontName="Times-Bold", fontSize=11.5, leading=14, spaceAfter=6)

    doc = SimpleDocTemplate(
        str(destination),
        pagesize=LETTER,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )

    story = [Paragraph(escape(title), h1), Spacer(1, 0.1 * inch)]

    idx = 0
    in_code = False
    code_lines: list[str] = []
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if idx == 0 and stripped.startswith("#"):
            idx += 1
            continue

        if stripped.startswith("```"):
            if in_code and code_lines:
                story.append(Preformatted("\n".join(code_lines), code))
                story.append(Spacer(1, 0.05 * inch))
                code_lines = []
            in_code = not in_code
            idx += 1
            continue

        if in_code:
            code_lines.append(raw.rstrip())
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        table_rows, next_idx = parse_table(lines, idx)
        if table_rows:
            pdf_table = PdfTable([[strip_inline_code(cell) for cell in row] for row in table_rows], repeatRows=1)
            pdf_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                        ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("LEADING", (0, 0), (-1, -1), 11),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(pdf_table)
            story.append(Spacer(1, 0.12 * inch))
            idx = next_idx
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            heading_text = stripped[level:].strip()
            style = h1 if level == 1 else h2 if level == 2 else h3
            story.append(Paragraph(escape(heading_text), style))
            idx += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            story.append(Paragraph(f"{numbered.group(1)}. {rich_text(numbered.group(2))}", body))
            idx += 1
            continue

        if stripped.startswith("- "):
            story.append(Paragraph(f"&#8226; {rich_text(stripped[2:].strip())}", body))
            idx += 1
            continue

        story.append(Paragraph(rich_text(stripped), body))
        idx += 1

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def sync_markdown_files() -> None:
    for rel_path in SYNC_FILES:
        src = ROOT / rel_path
        dst = BUNDLE_DIR / Path(rel_path).name
        shutil.copy2(src, dst)


def write_upload_manifest() -> None:
    manifest = """EMSE upload-set manifest

Recommended upload files:
- EMSE_main_manuscript.docx
- EMSE_title_page_and_declarations.docx
- EMSE_cover_letter.docx
- ESM_1.pdf

Supporting note:
- EMSE_supplement_source.docx is retained as the editable source used to generate ESM_1.pdf.
"""
    (UPLOAD_DIR / "README.txt").write_text(manifest, encoding="utf-8")


def main() -> None:
    BUNDLE_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    sync_markdown_files()
    for rel_src, out_name in REVIEW_DOCS:
        src = ROOT / rel_src
        dst = BUNDLE_DIR / out_name
        render_markdown(src, dst, branded=True, footer_text="Quantyra submission draft", include_page_numbers=False)
    for rel_src, out_name in UPLOAD_DOCS:
        src = ROOT / rel_src
        dst = UPLOAD_DIR / out_name
        render_markdown(
            src,
            dst,
            branded=False,
            footer_text=None,
            include_page_numbers=out_name == "EMSE_main_manuscript.docx",
        )
    render_markdown_pdf(ROOT / "docs/evidence/manuscript-supplement-v1.md", UPLOAD_DIR / "ESM_1.pdf")
    write_upload_manifest()


if __name__ == "__main__":
    main()
